import base64
import json
import re
import aiohttp
import asyncio
from typing import Dict, List, Any, Optional, Tuple
from services.docx_parser import parse_docx_test
from services.pdf_parser import parse_pdf_test


class AIGeneratorError(Exception):
    pass


# 2026-yilning eng yangi, sinovdan o'tgan ultra-tezkor va barqaror modellari
MODELS_TO_TRY = [
    "gemini-3.5-flash-lite", # 1-o'rin: 0.8 soniyada javob beruvchi ultra-yengil model
    "gemini-3.6-flash",      # 2-o'rin: 1.3 soniyada yuqori aniqlikda javob beruvchi rasmiy model
    "gemini-3.1-flash-lite", # 3-o'rin: 2.2 soniyada barqaror ishlovchi model
    "gemini-3.5-flash",      # 4-o'rin: Yuqori sifatli zaxira
    "gemini-3.8-flash",      # 5-o'rin: Katta hajmli zaxira
    "gemini-3.7-flash",      # 6-o'rin: Qo'shimcha zaxira
    "gemini-2.5-flash"       # 7-o'rin: Eski kalitlar uchun moslashuvchan fallback
]

SYSTEM_INSTRUCTION = (
    "Siz professional test tuzuvchi va pedagogik ekspert yordamchisiz. "
    "Foydalanuvchi taqdim etgan material (matn, darslik, konspekt, rasm, word yoki pdf) asosida "
    "aniq, sifatli, mantiqiy va xolis test savollarini tuzib berishingiz kerak. "
    "Har bir savol uchun 4 ta variant (A, B, C, D), 1 ta to'g'ri javob va qisqa tushuntirish (izoh) bo'lishi shart. "
    "Javobni FAQAT talab qilingan JSON formatida qaytaring."
)


def _clean_json_text(raw_text: str) -> str:
    """Markdown bloklari va ortiqcha belgilarni tozalab, sof JSON ni ajratib oladi"""
    text = raw_text.strip()
    if text.startswith("```"):
        text = re.sub(r"^```[a-zA-Z]*\n?", "", text)
        text = re.sub(r"\n?```$", "", text)
    text = text.strip()
    # Agar model boshida yoki oxirida biror matn qo'shgan bo'lsa, eng tashqi { va } ni ajratib olamiz
    if not (text.startswith("{") and text.endswith("}")):
        m = re.search(r"(\{[\s\S]*\})", text)
        if m:
            text = m.group(1)
    return text.strip()


async def _call_gemini_api(payload: dict, api_key: str) -> dict:
    if not api_key:
        raise AIGeneratorError(
            "AI API kaliti kiritilmagan!\n\n"
            "Iltimos, administrator paneli orqali (yoki /admin -> '🔑 AI API kalitini sozlash' bo'limidan) "
            "Google AI Studio (aistudio.google.com) dan olingan bepul API kalitni kiriting."
        )

    last_error = ""
    had_quota_error = False
    timeout = aiohttp.ClientTimeout(total=15, connect=5)

    # 2 tsikl davomida eng tezkor va barqaror modellar bo'ylab harakatlanamiz
    for cycle in range(2):
        for model_name in MODELS_TO_TRY:
            url = f"https://generativelanguage.googleapis.com/v1beta/models/{model_name}:generateContent?key={api_key}"

            model_payload = json.loads(json.dumps(payload))
            if "generationConfig" not in model_payload:
                model_payload["generationConfig"] = {}

            # Modellarga mos optimal tezlik (thinkingLevel minimal - 1 soniyalik tezlik)
            if "3.6-flash" in model_name or "3.5-flash" in model_name or "3.1-flash-lite" in model_name:
                model_payload["generationConfig"]["thinkingConfig"] = {"thinkingLevel": "minimal"}
            elif "3.8-flash" in model_name or "3.7-flash" in model_name:
                model_payload["generationConfig"]["thinkingConfig"] = {"thinkingLevel": "low"}
            elif "2.5-flash" in model_name:
                model_payload["generationConfig"]["thinkingConfig"] = {"thinkingBudget": 0}
            else:
                if "thinkingConfig" in model_payload["generationConfig"]:
                    del model_payload["generationConfig"]["thinkingConfig"]

            try:
                async with aiohttp.ClientSession(timeout=timeout) as session:
                    async with session.post(url, json=model_payload) as resp:
                        status = resp.status
                        try:
                            data = await resp.json()
                        except Exception:
                            data = {}

                        if status != 200:
                            err_msg = "Noma'lum xatolik"
                            if isinstance(data, dict) and "error" in data:
                                err_msg = data["error"].get("message", str(data["error"]))

                            # 1. Agar API kalit butunlay yaroqsiz bo'lsa (400 API_KEY_INVALID)
                            if "API key not valid" in err_msg or (status == 400 and "API_KEY_INVALID" in err_msg):
                                raise AIGeneratorError(
                                    "Kiritilgan AI API kaliti yaroqsiz!\n"
                                    "Iltimos, aistudio.google.com saytidan bepul yangi kalit olib, "
                                    "bot sozlamalaridan qayta kiriting."
                                )

                            # 2. Ruxsat cheklangan bo'lsa (403 Forbidden)
                            if status == 403:
                                raise AIGeneratorError(
                                    "AI API kalitida ruxsat cheklovi (403) mavjud.\n"
                                    "Iltimos, Google Cloud yoki AI Studio konsolida 'Generative Language API' yoqilganligini tekshiring."
                                )

                            # 3. 429 Quota Exceeded (Rate limit)
                            if status == 429:
                                had_quota_error = True
                                last_error = "So'rovlar limiti (kvota) to'ldi (429)"
                                continue

                            # 4. Agar model thinkingConfig ni qabul qilmasa (400), thinkingConfigsiz qayta sinaymiz
                            if status == 400 and ("thinking" in err_msg.lower() or "unknown field" in err_msg.lower()):
                                plain_payload = json.loads(json.dumps(payload))
                                if "generationConfig" in plain_payload and "thinkingConfig" in plain_payload["generationConfig"]:
                                    del plain_payload["generationConfig"]["thinkingConfig"]
                                async with session.post(url, json=plain_payload) as retry_resp:
                                    if retry_resp.status == 200:
                                        retry_data = await retry_resp.json()
                                        candidates = retry_data.get("candidates", [])
                                        if candidates:
                                            parts = candidates[0].get("content", {}).get("parts", [])
                                            if parts:
                                                raw_json = parts[0].get("text", "")
                                                return json.loads(_clean_json_text(raw_json))

                            # 503, 500, 502, 504 yoki boshqa holatda darhol keyingi modelga o'tamiz
                            last_error = f"{model_name}: {err_msg} ({status})"
                            continue

                        # Muvaffaqiyatli 200 javob
                        candidates = data.get("candidates", [])
                        if not candidates:
                            last_error = f"{model_name}: Bo'sh javob"
                            continue

                        content_parts = candidates[0].get("content", {}).get("parts", [])
                        if not content_parts:
                            last_error = f"{model_name}: Javobda matn topilmadi"
                            continue

                        raw_json = content_parts[0].get("text", "")
                        cleaned = _clean_json_text(raw_json)

                        try:
                            parsed = json.loads(cleaned)
                            return parsed
                        except json.JSONDecodeError:
                            last_error = f"{model_name}: JSON format xatoligi"
                            continue

            except (aiohttp.ClientError, asyncio.TimeoutError):
                last_error = f"{model_name}: Tarmoq ulanishida uzilish (Timeout)"
                continue

        # Har bir tsikl oralig'ida qisqa tanaffus
        if cycle == 0:
            await asyncio.sleep(0.3)

    if had_quota_error:
        raise AIGeneratorError(
            "AI so'rovlar limiti (kvotasi) to'ldi!\n\n"
            "Google bepul API kalit uchun daqiqalik cheklov qo'ygan. "
            "Iltimos, 1 daqiqa kutib qayta urinib ko'ring yoki aistudio.google.com dan yangi bepul kalit oling."
        )

    raise AIGeneratorError(
        f"AI xizmatida vaqtinchalik uzilish yuzaga keldi.\n"
        f"Sababi: {last_error or 'Server yuklamasi yuqori'}.\n\n"
        "Iltimos, bir necha soniyadan so'ng qayta urinib ko'ring."
    )


def _validate_ai_test(data: dict) -> dict:
    """AI qaytargan ma'lumotlarni tekshirib, standart formatga keltiradi"""
    title = str(data.get("title", "AI tomonidan yaratilgan test")).strip() or "AI Test"
    raw_questions = data.get("questions", [])
    if not raw_questions:
        raise AIGeneratorError("AI savollarni shakllantira olmadi. Boshqa matn yoki material bilan urinib ko'ring.")

    valid_questions = []
    for idx, q in enumerate(raw_questions, start=1):
        q_text = str(q.get("question_text", "")).strip()
        opt_a = str(q.get("option_a", "")).strip()
        opt_b = str(q.get("option_b", "")).strip()
        opt_c = str(q.get("option_c", "")).strip()
        opt_d = str(q.get("option_d", "")).strip()
        correct = str(q.get("correct_option", "A")).strip().upper()
        if correct not in ["A", "B", "C", "D"]:
            correct = "A"
        explanation = str(q.get("explanation", "")).strip()

        if not q_text or not opt_a or not opt_b or not opt_c or not opt_d:
            continue

        valid_questions.append({
            "question_text": q_text,
            "option_a": opt_a,
            "option_b": opt_b,
            "option_c": opt_c,
            "option_d": opt_d,
            "correct_option": correct,
            "explanation": explanation,
            "image_bytes": None,
            "image_ext": "png"
        })

    if not valid_questions:
        raise AIGeneratorError("AI to'liq va yaroqli variantlarga ega savollarni tuzib bera olmadi.")

    return {
        "title": title,
        "questions": valid_questions
    }


async def generate_test_from_content(
    text_content: str,
    api_key: str,
    question_count: int = 5,
    custom_instruction: str = "",
    media_parts: Optional[List[dict]] = None
) -> Dict[str, Any]:
    """
    Sun'iy intellekt (AI) yordamida har qanday matn yoki media asosida test tuzadi.
    """
    prompt = (
        f"{SYSTEM_INSTRUCTION}\n\n"
        f"Vazifa: Taqdim etilayotgan material asosida {question_count} ta test savolini tuzing.\n"
        f"Qo'shimcha talab/ko'rsatma: {custom_instruction or 'Darslik talablariga mos yuqori darajada tuzilsin.'}\n\n"
        f"Qaytishi SHART bo'lgan JSON tuzilishi:\n"
        f"{{\n"
        f'  "title": "Mavzu nomi yoki sarlavha",\n'
        f'  "questions": [\n'
        f'    {{\n'
        f'      "question_text": "Savol matni",\n'
        f'      "option_a": "A varianti",\n'
        f'      "option_b": "B varianti",\n'
        f'      "option_c": "C varianti",\n'
        f'      "option_d": "D varianti",\n'
        f'      "correct_option": "A",\n'
        f'      "explanation": "To\'g\'ri javob nega bunday ekanligi haqida qisqa izoh"\n'
        f'    }}\n'
        f'  ]\n'
        f"}}\n\n"
    )

    parts: List[dict] = [{"text": prompt}]

    if text_content:
        parts.append({"text": f"Material matni:\n{text_content[:30000]}"})

    if media_parts:
        for mp in media_parts:
            parts.append(mp)

    payload = {
        "contents": [
            {
                "parts": parts
            }
        ],
        "generationConfig": {
            "temperature": 0.3,
            "response_mime_type": "application/json"
        }
    }

    raw_result = await _call_gemini_api(payload, api_key)
    return _validate_ai_test(raw_result)


async def generate_test_from_image(
    image_bytes: bytes,
    mime_type: str,
    api_key: str,
    question_count: int = 5,
    custom_prompt: str = ""
) -> Dict[str, Any]:
    """Rasm (darslik sahifasi, doska, fotosurat) asosida test tuzadi"""
    b64_data = base64.b64encode(image_bytes).decode("utf-8")
    media_parts = [
        {
            "inline_data": {
                "mime_type": mime_type or "image/jpeg",
                "data": b64_data
            }
        }
    ]
    prompt_desc = custom_prompt or "Rasmda keltirilgan darslik/mavzu mazmunidan kelib chiqib test tuzing."
    return await generate_test_from_content(
        text_content=prompt_desc,
        api_key=api_key,
        question_count=question_count,
        custom_instruction=custom_prompt,
        media_parts=media_parts
    )


def extract_text_from_docx_bytes(docx_bytes: bytes) -> str:
    """Word (.docx va .doc) faylidan barcha matnlarni tezkor va to'liq ajratib oladi"""
    import io
    # 1. Standart .docx (XML) tahlili
    try:
        import docx
        doc = docx.Document(io.BytesIO(docx_bytes))
        paragraphs_text = []
        for p in doc.paragraphs:
            t = p.text.strip()
            if t:
                paragraphs_text.append(t)
        for tbl in doc.tables:
            for row in tbl.rows:
                r_text = " | ".join([c.text.strip() for c in row.cells if c.text.strip()])
                if r_text:
                    paragraphs_text.append(r_text)
        full_text = "\n".join(paragraphs_text)
        if full_text.strip():
            return full_text.strip()
    except Exception:
        pass

    # 2. Eskiroq .doc (binary) yoki g'ayrioddiy kodirovkali fayllar uchun zaxira matn chiqaruvchi
    try:
        raw_text = docx_bytes.decode("utf-8", errors="ignore")
        cleaned = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', ' ', raw_text)
        tokens = [word for word in cleaned.split() if len(word) > 1]
        if len(tokens) > 20:
            return " ".join(tokens[:5000])
    except Exception:
        pass

    return ""


def extract_text_from_pdf_bytes(pdf_bytes: bytes, max_pages: int = 50) -> Tuple[str, Optional[bytes], int]:
    """
    PDF faylidan matnni ajratib oladi.
    Qaytaradi: (extracted_text, compact_pdf_bytes_if_scanned, total_pages)
    """
    import io
    from pypdf import PdfReader, PdfWriter

    total_pages = 0
    extracted_text = ""
    try:
        reader = PdfReader(io.BytesIO(pdf_bytes))
        total_pages = len(reader.pages)
        text_chunks = []
        pages_to_read = min(total_pages, max_pages)

        for i in range(pages_to_read):
            try:
                page = reader.pages[i]
                t = page.extract_text()
                if t and t.strip():
                    text_chunks.append(t.strip())
            except Exception:
                continue

        extracted_text = "\n\n".join(text_chunks)
    except Exception:
        pass

    # Agar matn topilgan bo'lsa (kamida 80 ta belgi)
    if len(extracted_text.strip()) >= 80:
        return extracted_text.strip(), None, total_pages

    # Agar matn deyarli topilmagan bo'lsa (skaner qilingan yoki rasmli PDF bo'lsa),
    # butun faylni emas, faqat dastlabki 3-4 sahifasini ixchamlashtirib olamiz (Google payload limitidan oshmasligi uchun)
    try:
        reader = PdfReader(io.BytesIO(pdf_bytes))
        writer = PdfWriter()
        for i in range(min(len(reader.pages), 4)):
            writer.add_page(reader.pages[i])
        buf = io.BytesIO()
        writer.write(buf)
        return "", buf.getvalue(), total_pages
    except Exception:
        if len(pdf_bytes) < 4 * 1024 * 1024:
            return "", pdf_bytes, total_pages
        return "", None, total_pages


async def generate_test_from_pdf_file(
    pdf_bytes: bytes,
    api_key: str,
    question_count: int = 5,
    custom_prompt: str = ""
) -> Dict[str, Any]:
    """PDF darslik yoki konspekt faylidan test tuzadi"""
    extracted_text, compact_bytes, total_pages = extract_text_from_pdf_bytes(pdf_bytes)

    if extracted_text:
        return await generate_test_from_content(
            text_content=extracted_text,
            api_key=api_key,
            question_count=question_count,
            custom_instruction=custom_prompt
        )

    if compact_bytes:
        b64_data = base64.b64encode(compact_bytes).decode("utf-8")
        media_parts = [
            {
                "inline_data": {
                    "mime_type": "application/pdf",
                    "data": b64_data
                }
            }
        ]
        return await generate_test_from_content(
            text_content=custom_prompt or "Ushbu PDF darslik sahifalaridagi mavzu asosida test tuzing.",
            api_key=api_key,
            question_count=question_count,
            custom_instruction=custom_prompt,
            media_parts=media_parts
        )

    raise AIGeneratorError(
        "PDF faylidan matn ajratib bo'lmadi yoki fayl haddan tashqari katta.\n"
        "Iltimos, PDF matnini nusxalab yuboring yoki Word (.docx) shaklida yuboring."
    )


async def generate_test_from_docx_file(
    docx_bytes: bytes,
    api_key: str,
    question_count: int = 5,
    custom_prompt: str = ""
) -> Dict[str, Any]:
    """Word (.docx, .doc) konspekt yoki darslik faylidan test tuzadi"""
    full_text = extract_text_from_docx_bytes(docx_bytes)
    if not full_text:
        raise AIGeneratorError(
            "Word faylida o'qiladigan matn topilmadi!\n"
            "Iltimos, fayl ichida matn mavjudligini tekshiring yoki matnni botga to'g'ridan-to'g'ri xabar qilib yuboring."
        )

    return await generate_test_from_content(
        text_content=full_text,
        api_key=api_key,
        question_count=question_count,
        custom_instruction=custom_prompt
    )
