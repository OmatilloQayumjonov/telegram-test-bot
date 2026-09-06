import re
import os
from typing import Dict, List, Any, Optional, Tuple
import docx


class DocxParseError(Exception):
    pass


CYRILLIC_TO_LATIN_OPT = {
    'а': 'a', 'б': 'b', 'в': 'c', 'с': 'c', 'г': 'd', 'д': 'd',
    'А': 'A', 'Б': 'B', 'В': 'C', 'С': 'C', 'Г': 'D', 'Д': 'D'
}


def normalize_opt_letter(char: str) -> str:
    char_clean = char.strip()
    return CYRILLIC_TO_LATIN_OPT.get(char_clean, char_clean).upper()


def extract_images_from_element(element, doc) -> List[Tuple[bytes, str]]:
    """
    Word XML elementidan (Paragraph, Table yoki Cell) barcha rasmlarni (bytes, ext) ko'rinishida chiqaradi.
    """
    images = []
    try:
        for desc in element.iter():
            for attr_name, attr_val in desc.attrib.items():
                if attr_name.endswith('embed') or attr_name.endswith('id'):
                    if attr_val in doc.part.related_parts:
                        rel_part = doc.part.related_parts[attr_val]
                        content_type = getattr(rel_part, 'content_type', '')
                        if 'image' in content_type:
                            blob = getattr(rel_part, 'blob', None)
                            if blob:
                                ext = 'png'
                                if 'jpeg' in content_type or 'jpg' in content_type:
                                    ext = 'jpg'
                                elif 'gif' in content_type:
                                    ext = 'gif'
                                elif 'webp' in content_type:
                                    ext = 'webp'
                                elif hasattr(rel_part, 'partname'):
                                    ext = str(rel_part.partname).rsplit('.', 1)[-1].lower()
                                images.append((blob, ext))
    except Exception:
        pass
    return images


def format_table_to_text(table: docx.table.Table) -> str:
    """
    Savol ichidagi jadvalni Telegramda yaxshi ko'rinadigan monospace formatga aylantiradi.
    """
    rows_data = []
    for row in table.rows:
        row_cells = []
        seen_tcs = set()
        for cell in row.cells:
            if cell._tc in seen_tcs:
                continue
            seen_tcs.add(cell._tc)
            c_text = " ".join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
            row_cells.append(c_text)
        if any(row_cells):
            rows_data.append(row_cells)

    if not rows_data:
        return ""

    max_cols = max(len(r) for r in rows_data)
    for r in rows_data:
        while len(r) < max_cols:
            r.append("")

    # Ustunlar maksimal kengligi (Telegram ekrani uchun moslashtirilgan)
    col_widths = [max(len(r[c]) for r in rows_data) for c in range(max_cols)]
    col_widths = [min(max(w, 2), 30) for w in col_widths]

    sep = "+" + "+".join(["-" * (w + 2) for w in col_widths]) + "+"
    lines = ["📊 <b>Jadval:</b>", "<pre>", sep]

    for idx, r in enumerate(rows_data):
        row_str = "|" + "|".join([f" {r[col][:col_widths[col]].ljust(col_widths[col])} " for col in range(max_cols)]) + "|"
        lines.append(row_str)
        if idx == 0:
            lines.append(sep)

    lines.append(sep)
    lines.append("</pre>")
    return "\n".join(lines)


def is_questions_table(table: docx.table.Table) -> bool:
    """
    Jadval test savollarini joylashtirish uchun foydalanilganmi yoki savol ichidagi ma'lumot jadvalimi?
    """
    total_text = ""
    for row in table.rows:
        for cell in row.cells:
            total_text += " " + cell.text

    # Variant belgilari (A), B), C), D) yoki *A))
    opt_matches = re.findall(r"(?:^|\s)(?:[\*\+]?\s*)?[a-dA-Dа-яА-Я][\.\)]\s+", total_text)
    if len(opt_matches) >= 3:
        return True

    # "Javob:" yoki "To'g'ri javob:" yoki "Kalit:"
    if re.search(r"(?:javob|kalit|to['`‘]?g['`‘]?ri javob|answer)\s*:", total_text, re.IGNORECASE):
        return True

    # Sarlavha qatorida "Savol", "Variant", "Javob"
    if table.rows:
        first_row_text = " ".join([c.text.lower() for c in table.rows[0].cells])
        if "savol" in first_row_text and ("javob" in first_row_text or "a" in first_row_text):
            return True

    return False


def split_multi_options_line(line: str) -> List[str]:
    """
    Bitta qatorda bir nechta variant bo'lsa (A) 10  B) 20  C) 30  D) 40),
    ularni alohida qatorlarga ajratadi.
    """
    pattern = r"(?<=\S)\s{2,}(?=(?:[\*\+]?\s*)?[a-dA-Dа-яА-Я][\.\)\:\-])"
    parts = re.split(pattern, line)
    if len(parts) >= 2:
        return [p.strip() for p in parts if p.strip()]

    pattern2 = r"(?<=\S)\s+(?=(?:[\*\+]?\s*)?[b-dB-Dб-гБ-Г][\.\)\:\-])"
    parts2 = re.split(pattern2, line)
    if len(parts2) >= 2:
        return [p.strip() for p in parts2 if p.strip()]

    return [line]


def parse_docx_test(file_path_or_bytes, default_title: str = "Yangi test") -> Dict[str, Any]:
    """
    Word (.docx) faylidan rasmli, jadvalli, standart va noan'anaviy formatdagi savollarni tahlil qiladi.
    """
    try:
        doc = docx.Document(file_path_or_bytes)
    except Exception as e:
        raise DocxParseError(f"Word faylini ochishda xatolik: {str(e)}")

    doc_elements = []

    for child in doc.element.body:
        if child.tag.endswith('p'):
            p = docx.text.paragraph.Paragraph(child, doc)
            text = p.text.strip()
            images = extract_images_from_element(child, doc)
            doc_elements.append({
                "type": "paragraph",
                "text": text,
                "images": images
            })
        elif child.tag.endswith('tbl'):
            table = docx.table.Table(child, doc)
            images = extract_images_from_element(child, doc)

            if is_questions_table(table):
                # Savollar joylashtirilgan jadval
                for row in table.rows:
                    for cell in row.cells:
                        cell_images = extract_images_from_element(cell._tc, doc)
                        for p in cell.paragraphs:
                            c_text = p.text.strip()
                            p_imgs = extract_images_from_element(p._element, doc)
                            all_imgs = p_imgs if p_imgs else cell_images
                            if c_text or all_imgs:
                                doc_elements.append({
                                    "type": "paragraph",
                                    "text": c_text,
                                    "images": all_imgs
                                })
            else:
                # Savol ichidagi ma'lumot jadvali (data table)
                table_text = format_table_to_text(table)
                doc_elements.append({
                    "type": "table_data",
                    "text": table_text,
                    "images": images
                })

    if not doc_elements:
        raise DocxParseError("Word fayli bo'sh yoki unda ma'lumot topilmadi.")

    title = default_title
    # Sarlavhani topish (birinchi qatorda test nomi bo'lsa)
    for i, elem in enumerate(doc_elements):
        if elem["type"] == "paragraph" and elem["text"]:
            first_line = elem["text"]
            if not re.match(r"^(?:№\s*)?(?:(?:savol|question|вопрос)\s*)?(\d+)", first_line, re.IGNORECASE):
                if not re.match(r"^([\*\+]?\s*)?[a-dA-Dа-яА-Я][\.\)]", first_line):
                    title = first_line
                    doc_elements.pop(i)
            break

    questions: List[Dict[str, Any]] = []
    current_q = None
    pending_image = None
    pending_table = None

    option_pattern = re.compile(r"^([\*\+])?\s*([a-dA-Dа-яА-Я])[\.\)\:\-]\s*(.*)$", re.IGNORECASE)
    ans_pattern = re.compile(r"^(?:javob|to['`‘]?g['`‘]?ri javob|kalit|answer|correct(?:\s*answer)?|жавоб|тўғри жавоб|калит)\s*[:\-=]\s*\(?([a-dA-Dа-яА-Я])", re.IGNORECASE)
    explanation_pattern = re.compile(r"^(?:izoh|tushuntirish|sharh|explanation|изоҳ|шарҳ)\s*[:\-=]\s*(.*)$", re.IGNORECASE)
    q_start_pattern = re.compile(r"^(?:№\s*)?(?:(?:savol|question|вопрос)\s*)?(\d+)(?:[\.\)\-:\s]+(?:savol|question|вопрос)?)?[\.\)\-:\s]*(.*)$", re.IGNORECASE)

    def finalize_question(q):
        if not q:
            return

        # Agar savol matni bo'sh bo'lsa, lekin rasm bo'lsa
        if not q['text'].strip():
            if q.get('image_bytes'):
                q['text'] = "Quyidagi rasmga asosan to'g'ri javobni tanlang:"
            else:
                q['text'] = f"{q['number']}-savol:"

        missing = []
        for opt in ['a', 'b', 'c', 'd']:
            if not q['options'].get(opt):
                missing.append(opt.upper())

        if missing:
            raise DocxParseError(
                f"'{q['number']}-savol' uchun quyidagi variantlar topilmadi: {', '.join(missing)}.\n"
                f"Savol matni: {q['text'][:60]}..."
            )

        if not q.get('correct_option'):
            raise DocxParseError(
                f"'{q['number']}-savol' uchun to'g'ri javob belgilanmagan.\n"
                f"Masalan: 'Javob: A' deb yozing yoki to'g'ri variant oldiga * qo'ying.\n"
                f"Savol matni: {q['text'][:60]}..."
            )

        questions.append({
            "question_text": q['text'].strip(),
            "option_a": q['options']['a'].strip(),
            "option_b": q['options']['b'].strip(),
            "option_c": q['options']['c'].strip(),
            "option_d": q['options']['d'].strip(),
            "correct_option": q['correct_option'].upper(),
            "explanation": q.get('explanation', '').strip(),
            "image_bytes": q.get('image_bytes'),
            "image_ext": q.get('image_ext', 'png')
        })

    for elem in doc_elements:
        elem_type = elem["type"]
        images = elem.get("images", [])

        # Rasmlarni biriktirish
        if images:
            if current_q and not current_q.get("image_bytes"):
                current_q["image_bytes"], current_q["image_ext"] = images[0]
            elif not current_q and not pending_image:
                pending_image = images[0]

        if elem_type == "table_data":
            table_str = elem.get("text", "")
            if current_q:
                current_q['text'] = (current_q['text'] + "\n\n" + table_str).strip()
            else:
                pending_table = table_str
            continue

        raw_text = elem.get("text", "")
        if not raw_text:
            continue

        lines = raw_text.split("\n")
        expanded_lines = []
        for l in lines:
            l_clean = l.strip()
            if l_clean:
                expanded_lines.extend(split_multi_options_line(l_clean))

        for line in expanded_lines:
            line_clean = line.strip()
            if not line_clean:
                continue

            q_match = q_start_pattern.match(line_clean)
            opt_match = option_pattern.match(line_clean)
            ans_match = ans_pattern.match(line_clean)
            exp_match = explanation_pattern.match(line_clean)

            # 1. Yangi savol boshlanishi
            if q_match and not opt_match and not ans_match and not exp_match:
                if current_q:
                    finalize_question(current_q)

                q_num = q_match.group(1)
                q_text = q_match.group(2).strip()

                current_q = {
                    "number": q_num,
                    "text": q_text,
                    "options": {},
                    "correct_option": None,
                    "explanation": "",
                    "image_bytes": None,
                    "image_ext": "png"
                }

                if pending_image:
                    current_q["image_bytes"], current_q["image_ext"] = pending_image
                    pending_image = None

                if pending_table:
                    current_q["text"] = (current_q["text"] + "\n\n" + pending_table).strip()
                    pending_table = None

                continue

            # 2. Variantlar (A) B) C) D) yoki *A) +B) yoki Kirill А) Б))
            if opt_match:
                if not current_q:
                    current_q = {
                        "number": str(len(questions) + 1),
                        "text": "Savol",
                        "options": {},
                        "correct_option": None,
                        "explanation": "",
                        "image_bytes": None,
                        "image_ext": "png"
                    }
                    if pending_image:
                        current_q["image_bytes"], current_q["image_ext"] = pending_image
                        pending_image = None
                    if pending_table:
                        current_q["text"] = (current_q["text"] + "\n\n" + pending_table).strip()
                        pending_table = None

                mark = opt_match.group(1)
                raw_letter = opt_match.group(2)
                letter = normalize_opt_letter(raw_letter).lower()
                val = opt_match.group(3).strip()

                current_q['options'][letter] = val
                if mark in ['*', '+']:
                    current_q['correct_option'] = letter.upper()
                continue

            # 3. Javob kaliti (Javob: B)
            if ans_match:
                if current_q:
                    raw_ans = ans_match.group(1)
                    current_q['correct_option'] = normalize_opt_letter(raw_ans)
                continue

            # 4. Izoh
            if exp_match:
                if current_q:
                    current_q['explanation'] = exp_match.group(1).strip()
                continue

            # 5. Savol matnining davomi (ko'p qatorli savollar uchun)
            if current_q and not current_q['options']:
                if current_q['text']:
                    current_q['text'] += "\n" + line_clean
                else:
                    current_q['text'] = line_clean
            elif current_q and current_q.get('explanation'):
                current_q['explanation'] += " " + line_clean

    if current_q:
        finalize_question(current_q)

    if not questions:
        raise DocxParseError(
            "Fayldan birorta ham to'g'ri formatdagi savol topilmadi.\n"
            "Iltimos, har bir savol raqami (1.), variantlar (A, B, C, D) va to'g'ri javob ko'rsatilganiga ishonch hosil qiling."
        )

    return {
        "title": title,
        "questions": questions
    }
