import docx
from pathlib import Path


def create_sample_docx(output_path: str = "namuna.docx") -> str:
    """O'qituvchi uchun namunaviy test Word faylini yaratadi"""
    doc = docx.Document()

    # Sarlavha
    doc.add_heading("Informatika va Axborot Texnologiyalari: 1-Mavzu Testi", level=1)

    p_intro = doc.add_paragraph(
        "Ushbu fayl bot uchun namunaviy test fayli hisoblanadi. Savollarni quyidagi tartibda yozishingiz mumkin:\n"
    )

    # 1-savol: Standart format (Javob va Izoh bilan)
    doc.add_paragraph("1. Python dasturlash tili qaysi yilda yaratilgan?")
    doc.add_paragraph("A) 1989-yil")
    doc.add_paragraph("B) 1991-yil")
    doc.add_paragraph("C) 1995-yil")
    doc.add_paragraph("D) 2000-yil")
    doc.add_paragraph("Javob: B")
    doc.add_paragraph("Izoh: Python dasturlash tili 1991-yilda Gvido van Rossum tomonidan ommaga taqdim etilgan.")
    doc.add_paragraph("")

    # 2-savol: Yulduzcha (*) bilan to'g'ri javobni belgilash
    doc.add_paragraph("2. Kompyuterning 'miyasi' deb ataluvchi asosiy qurilma nima?")
    doc.add_paragraph("A) Qattiq disk (HDD/SSD)")
    doc.add_paragraph("*B) Markaziy protsessor (CPU)")
    doc.add_paragraph("C) Tezkor xotira (RAM)")
    doc.add_paragraph("D) Ona plata (Motherboard)")
    doc.add_paragraph("Izoh: Protsessor (CPU) barcha hisoblash va mantiqiy amallarni bajaruvchi asosiy qismdir.")
    doc.add_paragraph("")

    # 3-savol: Izohsiz oddiy test
    doc.add_paragraph("3. Quyidagilardan qaysi biri operatsion tizim emas?")
    doc.add_paragraph("A) Windows 11")
    doc.add_paragraph("B) Ubuntu Linux")
    doc.add_paragraph("C) macOS")
    doc.add_paragraph("D) Google Chrome")
    doc.add_paragraph("To'g'ri javob: D")
    doc.add_paragraph("Izoh: Google Chrome operatsion tizim emas, balki veb-brauzer hisoblanadi.")
    doc.add_paragraph("")

    # 4-savol: Ko'p qatorli savol
    doc.add_paragraph("4. Algoritmning asosiy xossalaridan biri bu cheklilik (diskretlik) xossasidir.\nUshbu xossa nimani anglatadi?")
    doc.add_paragraph("A) Algoritm cheksiz takrorlanishi kerak")
    doc.add_paragraph("B) Algoritm aniq ketma-ketlikdagi chekli qadamlardan iborat bo'lishi kerak")
    doc.add_paragraph("C) Algoritm faqat bitta natija berishi kerak")
    doc.add_paragraph("D) Algoritm faqat kompyuterda ishlashi kerak")
    doc.add_paragraph("Javob: B")
    doc.add_paragraph("Izoh: Diskretlik - bu jarayonning alohida, chekli qadamlarga bo'linishini ifodalaydi.")
    doc.add_paragraph("")

    # 5-savol: Jadvalli savol namunasi
    doc.add_paragraph("5. Quyidagi jadval ma'lumotlariga qarab to'g'ri xulosani aniqlang:")
    tbl = doc.add_table(rows=3, cols=2)
    tbl.rows[0].cells[0].text = "Xotira turi"
    tbl.rows[0].cells[1].text = "Tezlik darajasi"
    tbl.rows[1].cells[0].text = "Kesh xotira (Cache)"
    tbl.rows[1].cells[1].text = "Juda yuqori"
    tbl.rows[2].cells[0].text = "Doimiy xotira (HDD)"
    tbl.rows[2].cells[1].text = "Past"
    doc.add_paragraph("A) HDD kesh xotiradan tezroq")
    doc.add_paragraph("*B) Kesh xotira eng yuqori tezlikka ega")
    doc.add_paragraph("C) Ikkala xotira tezligi teng")
    doc.add_paragraph("D) Barcha ma'lumotlar noto'g'ri")
    doc.add_paragraph("Izoh: Kesh xotira to'g'ridan-to'g'ri protsessor ichida joylashgan bo'lib, eng tezkor hisoblanadi.")

    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return str(path.resolve())


if __name__ == "__main__":
    create_sample_docx("namuna.docx")
    print("Namuna fayl yaratildi: namuna.docx")
